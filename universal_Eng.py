import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from io import BytesIO

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="Universal Lesson Planner", layout="wide")
st.title("🎓 Universal Lesson Planner generalisation")

# --- MAIN PAGE API KEY ENTRY (ABOVE TOPIC) ---
user_api_key = st.text_input(
    "🔑 Enter your own Gemini API Key:", 
    type="password", 
    help="Get your API key from Google AI Studio using your Gmail account."
)

# Helper function to dynamically check and load models based on the user's key
def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except Exception as e:
        st.error(f"Invalid API Key or Connection Error: {str(e)}")
        return None
    return "models/gemini-1.5-flash"  # Default fallback


# Process model assignment if the key is provided
selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)
    if selected_model_name:
        st.info(f"System connected via your API key. Active Model: {selected_model_name}")
else:
    st.warning("⚠️ Please enter your personal Gemini API Key above to start.")


# --- 2. AI LOGIC (INTEGRATED CRITERIA) ---
def generate_advanced_plan(topic, syllabus, extra_context, api_key, model_name):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    Topic: {topic}. Syllabus Code: {syllabus}. Context: {extra_context}.
    Generate a professional lesson plan in English.
    
    Use the following EXACT markers for the document structure:
    
    SECTION: TOPIC
    [Please display the topic here following the input parameter from user]
    SECTION: LESSON OBJECTIVES
    [4 points]
    SECTION: LESSON OUTCOMES
    [4 points]
    SECTION: SUCCESS CRITERIA
    [4 points]
    SECTION: PREREQUISITE
    [1 point]
    SECTION: KEYWORDS
    [6 items]
    SECTION: HOTS
    [4 main domains from Bloom's Taxonomy]
    SECTION: DIGITAL CITIZENSHIP
    [4 points on ethical tech use/Chromebooks/Canva/YouTube]

    SECTION: OPENING LESSON CONTENT
    [Hook activity and transition plan]

    SECTION: DIFFERENTIATION STRATEGIES (GREEN)
    - HA (Higher Achiever): [1 challenging activity]

    SECTION: DIFFERENTIATION STRATEGIES (YELLOW)
    - MA (Medium Achiever): [1 core activity]

    SECTION: DIFFERENTIATION STRATEGIES (RED)
    - LA (Lower Achiever): [1 scaffolded activity]

    SECTION: BLENDED LEARNING Activity ONE (15 MINS)
    - Activity 1: [Descriptions]
    - ----------------------------------------------------------------------------
    - Teacher Preparation: [Step-by-step before lesson]
    - ----------------------------------------------------------------------------
    - Objectives: [3 points]
    - ----------------------------------------------------------------------------
    - Student Tasks: [Step-by-step details]

    SECTION: BLENDED LEARNING Activity TWO (15 MINS)
    - Activity 2: [Descriptions]
    - -----------------------------------------------------------------------------
    - Teacher Preparation: [Step-by-step before lesson]
    - -----------------------------------------------------------------------------
    - Objectives: [3 points]
    - -----------------------------------------------------------------------------
    - Student Tasks: [Step-by-step details]
    
    SECTION: PLENARY (EXIT TICKET)
    [2-3 minute closing activity]

    SECTION: HOMEWORK
    [Task assigned based on topic]

    SECTION: SUGGESTED WAY FORWARD TASK
    - [Hook activity and transition plan for next day lesson]
    
    """
    try:
        response = model.generate_content(prompt)
        
        # Safety check for empty or interrupted candidate content responses
        if response.candidates and response.candidates[0].content.parts:
            return response.text
        else:
            return "⚠️ The AI returned an empty response. You might be clicking too fast and hitting your free API minute limit (15 requests/min). Please wait 60 seconds and try again."
            
    except Exception as e:
        return f"System Error: {str(e)}"

# --- 3. WORD EXPORT LOGIC ---
def create_word_export(topic, syllabus, text):
    doc = Document()
    doc.add_heading(f'PTES Lesson Plan: {topic}', 0)

    # Admin Header
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Week No:", "Date:"], ["Class Size:", "Day:"], ["Venue:", "Duration:"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    doc.add_paragraph()

    # Parsing and Boxing ALL Sections
    sections = text.split('SECTION:')
 
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        title = lines[0].strip()
        
        content = "\n".join(lines[1:]).strip().replace("**", "") 
        
        doc.add_heading(title.title(), level=1)
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = content
        doc.add_paragraph()
     
    # HOD Approval
    doc.add_page_break()
    doc.add_heading("HOD Approval & Remarks", level=1)
    hod_table = doc.add_table(rows=2, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Remarks:"
    hod_table.rows[1].height = Pt(50)
    hod_table.cell(1, 0).text = "Date:"; hod_table.cell(1, 1).text = "Signature:"

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. GUI ---
st.write("---") # Divider row under API bar
st.info("Type in the lesson topic, the subject's syllabus code and the extra information like canva, youtube, infographic")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Lesson Topic:")
with c2: u_syllabus = st.text_input("Syllabus Code:")
u_extra = st.text_area("Extra Context (Optional):")

if st.button("🚀 GENERATE COMPLETE LESSON PLAN"):
    if not user_api_key:
        st.error("❌ Please input your Google Gemini API key at the top before clicking generate.")
    elif not u_topic or not u_syllabus:
        st.error("❌ Please fill in the Topic and Syllabus fields.")
    else:
        with st.spinner("AI is integrating all criteria into your plan..."):
            result = generate_advanced_plan(u_topic, u_syllabus, u_extra, user_api_key, selected_model_name)
            st.session_state['adv_plan_out'] = result

if 'adv_plan_out' in st.session_state:
    st.divider()
    st.subheader("AI Draft Preview")
    st.text_area("Content", st.session_state['adv_plan_out'], height=400)
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['adv_plan_out'])
    st.download_button("📥 Download to Word version (.docx)", doc_file, f"Universal_LP_{u_topic}.docx")

st.markdown("---")
st.caption("Lesson planner 3.0 | Developer: Hjh Nurul Haziqah Hj Nordin | © 2026 PTES Innovation")
